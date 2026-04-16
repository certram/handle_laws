import logging
import re
from datetime import date
from pathlib import Path

from docxtpl import DocxTemplate

from ai_extractor import extract_bank_name as _ai_extract_bank_name
from config import CIVIL_FIRST
from ai_extractor import extract_alipay_account as _ai_extract_alipay
from ai_extractor import extract_tenpay_account as _ai_extract_tenpay
from ai_extractor import extract_property_info as _ai_extract_property

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent / "templates"

# 财产线索类型 → 模板文件名映射
CLUE_TYPE_MAP = {
    "银行": "bank",
    "银行账户": "bank",
    "支付宝": "alipay",
    "财付通": "tenpay",
    "房产": "property",
    "股权": "equity",
    "车辆": "vehicle",
}


def _extract_year(case_number: str) -> str:
    """从案号中提取年份，如 (2025)粤0305民初42226号 → 2025"""
    match = re.search(r"[(（](\d{4})[)）]", case_number)
    return match.group(1) if match else ""


def _extract_case_number(case_number: str) -> str:
    """从案号中提取纯数字部分，如 (2025)粤0305民初42226号 → 42226"""
    match = re.search(r"民初(\d+)号", case_number)
    return match.group(1) if match else case_number


def _get_template_name(clue_type: str) -> str | None:
    """根据财产线索类型返回模板文件名（不含扩展名）。不支持的类型返回 None。"""
    for keyword, template_name in CLUE_TYPE_MAP.items():
        if keyword in clue_type:
            return template_name
    return None


def _get_id_number(party: dict) -> str:
    """获取当事人身份号码（公司取信用代码，个人取身份证）。"""
    if party.get("类型") == "公司":
        return party.get("统一社会信用代码", "")
    return party.get("身份证号码", "")


def _get_id_label(party: dict) -> str:
    """获取身份号码的标签（公司→统一社会信用代码，个人→公民身份号码）。"""
    if party.get("类型") == "公司":
        return "统一社会信用代码"
    return "公民身份号码"


def _get_party_name(party: dict) -> str:
    """获取当事人名称。"""
    if party.get("类型") == "公司":
        return party.get("全称", "")
    return party.get("姓名", "")


def _join_party_names(parties: list[dict]) -> str:
    """将多个当事人名称用顿号拼接，如 'A公司、B公司'。"""
    names = [_get_party_name(p) for p in parties if _get_party_name(p)]
    return "、".join(names)


def _match_defendant_by_clue(defendants: list[dict], clue_detail: str) -> dict:
    """
    根据财产线索的详细内容匹配对应的被告。

    匹配策略：检查每个被告名称是否出现在财产线索文本中。
    """
    for defendant in defendants:
        name = _get_party_name(defendant)
        if name and name in clue_detail:
            return defendant
    # 无法匹配时取第一个
    return defendants[0] if defendants else {}


def _extract_bank_name(text: str) -> str:
    """从银行线索文本中用AI提取开户行名称。"""
    return _ai_extract_bank_name(text)


def _extract_account(text: str) -> str:
    """从银行线索文本中提取账号。"""
    # 格式1：账号：62608031256925 或 账户：6230580000198301488
    match = re.search(r"账[号户][：:]\s*(\d+)", text)
    if match:
        return match.group(1)
    # 格式2：751078594922（中国银行深圳锦绣支行）
    match = re.match(r"(\d+)[（(]", text.strip())
    if match:
        return match.group(1)
    return ""


def _build_context(case_data: dict, clue: dict) -> dict:
    """
    从案件数据和单条财产线索构建模板变量字典。

    模板变量说明：
    - plaintiff_name: 所有原告拼接（用于案件描述）
    - defendant_name: 所有被告拼接（用于案件描述）
    - target_name: 财产线索对应的被告名称（用于冻结/查封具体对象）
    - target_id: 财产线索对应的被告身份号码
    """
    case_info = case_data.get("案件基础信息", {})
    full_case_number = case_info.get("案号", "")
    year = _extract_year(full_case_number)
    case_number = _extract_case_number(full_case_number)

    # 所有原告拼接（案件描述用）
    plaintiffs = case_data.get("原告", [])
    plaintiff_name = _join_party_names(plaintiffs)

    # 所有被告拼接（案件描述用）
    defendants = case_data.get("被告", [])
    defendant_name = _join_party_names(defendants)

    # 根据财产线索匹配对应的被告（冻结/查封对象用）
    clue_detail = clue.get("详细内容", "")
    target_defendant = _match_defendant_by_clue(defendants, clue_detail)
    target_name = _get_party_name(target_defendant)
    target_id = _get_id_number(target_defendant)
    target_id_label = _get_id_label(target_defendant)

    # 保全金额
    preservation_info = case_data.get("保全信息", {})
    freeze_amount = preservation_info.get("申请保全总金额", "")

    # 通用变量
    context = {
        # 案件信息
        "year": year,
        "case_number": case_number,
        "full_case_number": full_case_number,
        "dispute_type": case_info.get("案由", ""),
        "freeze_amount": freeze_amount,
        "court_name": case_info.get("立案法院", ""),
        "judge_name": case_info.get("承办法官", ""),
        "court_phone": "",
        "court_address": "",
        # 案件描述：所有当事人
        "plaintiff_name": plaintiff_name,
        "defendant_name": defendant_name,
        # 冻结/查封具体对象：财产线索对应的被告
        "target_name": target_name,
        "target_id": target_id,
        "target_id_label": target_id_label,
    }

    # 类型专属变量（优先使用AI预提取的结构化字段，避免二次AI调用）
    clue_type = clue.get("类型", "")
    template_name = _get_template_name(clue_type)

    if template_name == "bank":
        bank_name = clue.get("开户银行", "")
        bank_account = clue.get("银行账号", "")
        # 如果预提取字段为空，回退到二次提取
        if not bank_name:
            bank_name = _extract_bank_name(clue_detail)
        if not bank_account:
            bank_account = _extract_account(clue_detail)
        context["bank_name"] = bank_name
        context["bank_account_number"] = bank_account

    elif template_name == "alipay":
        alipay_number = clue.get("支付宝账号", "")
        if not alipay_number:
            alipay_number = _ai_extract_alipay(clue_detail)
        # 验证：优先使用11位手机号
        if not re.match(r"^1\d{10}$", alipay_number):
            phone_match = re.search(r"1[3-9]\d{9}", clue_detail)
            if phone_match:
                logger.warning("支付宝账号非手机号，从原文正则补充: '%s' → '%s'", alipay_number, phone_match.group())
                alipay_number = phone_match.group()
        context["alipay_number"] = alipay_number

    elif template_name == "tenpay":
        tenpay_account = clue.get("微信号", "")
        if not tenpay_account:
            tenpay_account = _ai_extract_tenpay(clue_detail)
        context["account"] = tenpay_account

    elif template_name == "property":
        property_address = clue.get("房产地址", "")
        property_cert = clue.get("不动产证号", "")
        if not property_address or not property_cert:
            property_info = _ai_extract_property(clue_detail)
            if not property_address:
                property_address = property_info.get("property_address", "")
            if not property_cert:
                property_cert = property_info.get("property_certificate_number", "")
        context["property_address"] = property_address
        context["property_certificate_number"] = property_cert
        context["property_bureau"] = ""

    elif template_name == "equity":
        context["company_name"] = ""
        context["company_credit_code"] = ""
        context["share_percentage"] = ""
        context["share_capital"] = ""
        context["equity_bureau"] = ""

    elif template_name == "vehicle":
        # 从详细内容中提取车牌号（去掉"车牌号："前缀）
        plate_match = re.search(r"[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤川青藏琼宁][A-HJ-NP-Z][A-HJ-NP-Z0-9]{4,5}[A-HJ-NP-Z0-9挂学警港澳]", clue_detail)
        if plate_match:
            context["plate_number"] = plate_match.group()
        else:
            # 回退：去掉常见前缀
            plate = re.sub(r"^车牌号[：:]\s*", "", clue_detail).strip()
            context["plate_number"] = plate
        context["vehicle_bureau"] = ""

    return context


# ---------- 保全裁定书相关 ----------

_DIGIT_MAP = "〇一二三四五六七八九"


def _digit_to_chinese(d: int) -> str:
    """单个数字转中文：0→〇, 1→一, ..., 9→九。"""
    return _DIGIT_MAP[d]


def _year_to_chinese(year: int) -> str:
    """年份逐位转中文：2025 → 二〇二五。"""
    return "".join(_digit_to_chinese(int(d)) for d in str(year))


def _num_to_chinese(n: int) -> str:
    """1-31 的数字转中文：1→一, 10→十, 11→十一, 20→二十。"""
    if n <= 0:
        return ""
    if n < 10:
        return _digit_to_chinese(n)
    if n == 10:
        return "十"
    if n < 20:
        return "十" + _digit_to_chinese(n - 10)
    if n % 10 == 0:
        return _digit_to_chinese(n // 10) + "十"
    return _digit_to_chinese(n // 10) + "十" + _digit_to_chinese(n % 10)


def _date_to_chinese(d: date) -> str:
    """日期转中文格式：2025-08-18 → 二〇二五年八月十八日。"""
    return f"{_year_to_chinese(d.year)}年{_num_to_chinese(d.month)}月{_num_to_chinese(d.day)}日"


def _format_birth_date(date_str: str) -> str:
    """将 YYYY-MM-DD 格式转为 YYYY年M月D日。"""
    if not date_str:
        return ""
    parts = date_str.split("-")
    if len(parts) != 3:
        return date_str
    y, m, d = parts
    return f"{y}年{int(m)}月{int(d)}日"


def _format_party_for_ruling(party: dict, role: str) -> list[str]:
    """
    将当事人信息格式化为裁定书中的当事人段落文本列表。

    个人：["被告：易小兵，男，1973年3月4日出生，汉族，身份证住址XXX，公民身份号码XXX。"]
    公司：["被告：XXX公司，住所地XXX，统一社会信用代码XXX。", "法定代表人：XXX。"]
    """
    name = _get_party_name(party)
    if not name:
        return []

    if party.get("类型") == "公司":
        parts = [f"{role}：{name}"]
        addr = party.get("所在地", "")
        if addr:
            parts.append(f"住所地{addr}")
        credit_code = party.get("统一社会信用代码", "")
        if credit_code:
            parts.append(f"统一社会信用代码{credit_code}")
        lines = ["，".join(parts) + "。"]
        legal_rep = party.get("法定代表人", "")
        if legal_rep:
            lines.append(f"法定代表人：{legal_rep}。")
        return lines

    # 个人
    parts = [f"{role}：{name}"]
    gender = party.get("性别", "")
    if gender:
        parts.append(gender)
    birth = party.get("出生日期", "")
    if birth:
        parts.append(f"{_format_birth_date(birth)}出生")
    ethnicity = party.get("民族", "")
    if ethnicity:
        if not ethnicity.endswith("族"):
            ethnicity += "族"
        parts.append(ethnicity)
    addr = party.get("住址", "")
    if addr:
        parts.append(f"身份证住址{addr}")
    id_number = party.get("身份证号码", "")
    if id_number:
        parts.append(f"公民身份号码{id_number}")
    return ["，".join(parts) + "。"]


def _build_ruling_context(case_data: dict) -> dict:
    """构建保全裁定书的模板变量。"""
    case_info = case_data.get("案件基础信息", {})
    full_case_number = case_info.get("案号", "")
    year = _extract_year(full_case_number)
    case_number = _extract_case_number(full_case_number)

    plaintiffs = case_data.get("原告", [])
    defendants = case_data.get("被告", [])

    preservation_info = case_data.get("保全信息", {})
    guarantee_info = case_data.get("担保信息", {})

    # 将每个当事人展开为多行（公司含法定代表人另起一行）
    plaintiff_lines = []
    for p in plaintiffs:
        plaintiff_lines.extend(_format_party_for_ruling(p, "原告"))

    defendant_lines = []
    for d in defendants:
        defendant_lines.extend(_format_party_for_ruling(d, "被告"))

    return {
        "year": year,
        "case_number": case_number,
        "plaintiff_lines": plaintiff_lines,
        "defendant_lines": defendant_lines,
        "plaintiff_name": _join_party_names(plaintiffs),
        "defendant_name": _join_party_names(defendants),
        "dispute_type": case_info.get("案由", ""),
        "freeze_amount": preservation_info.get("申请保全总金额", ""),
        "guarantor_name": guarantee_info.get("担保人名称", ""),
        "guarantee_type": "信用",
        "ruling_date": _date_to_chinese(date.today()),
    }


def _replace_in_runs(paragraph, old_text: str, new_text: str):
    """在段落的 runs 中替换文本，保留第一个 run 的格式。"""
    full_text = "".join(r.text for r in paragraph.runs)
    if old_text not in full_text:
        return
    # 将所有 runs 的文本合并到第一个 run，清空其余
    new_full_text = full_text.replace(old_text, new_text)
    paragraph.runs[0].text = new_full_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _generate_ruling(
    case_data: dict, template_path: Path, output_dir: Path, generated_files: list[Path]
):
    """用 python-docx 直接生成保全裁定书（支持多行原告/被告段落）。"""
    from docx import Document

    context = _build_ruling_context(case_data)
    doc = Document(str(template_path))

    # 找到原告占位段落和被告占位段落
    plaintiff_para = None
    defendants_para = None
    for p in doc.paragraphs:
        if "{{ plaintiff_block }}" in p.text:
            plaintiff_para = p
        elif "{{ defendants_block }}" in p.text:
            defendants_para = p

    # 插入原告段落（公司可能有法定代表人另起一行）
    if plaintiff_para is not None:
        plaintiff_lines = context.get("plaintiff_lines", [])
        if plaintiff_lines:
            _replace_in_runs(plaintiff_para, "{{ plaintiff_block }}", plaintiff_lines[0])
            _insert_extra_paragraphs(plaintiff_para, plaintiff_lines[1:])

    # 插入被告段落（多个被告 + 公司法定代表人）
    if defendants_para is not None:
        defendant_lines = context.get("defendant_lines", [])
        if defendant_lines:
            _replace_in_runs(defendants_para, "{{ defendants_block }}", defendant_lines[0])
            _insert_extra_paragraphs(defendants_para, defendant_lines[1:])

    # 替换所有其他 {{ }} 变量
    simple_vars = {
        "{{ year }}": context.get("year", ""),
        "{{ case_number }}": context.get("case_number", ""),
        "{{ plaintiff_name }}": context.get("plaintiff_name", ""),
        "{{ defendant_name }}": context.get("defendant_name", ""),
        "{{ dispute_type }}": context.get("dispute_type", ""),
        "{{ freeze_amount }}": context.get("freeze_amount", ""),
        "{{ guarantor_name }}": context.get("guarantor_name", ""),
        "{{ guarantee_type }}": context.get("guarantee_type", ""),
        "{{ ruling_date }}": context.get("ruling_date", ""),
    }
    for p in doc.paragraphs:
        for old, new in simple_vars.items():
            if old in p.text:
                _replace_in_runs(p, old, new)

    case_number = context.get("case_number", "")
    today = date.today()
    date_str = f"{str(today.year)[2:]}{today.month:02d}{today.day:02d}"
    ruling_name = f"{CIVIL_FIRST}{case_number}保全裁定书（{date_str}）.docx"
    output_path = output_dir / ruling_name
    doc.save(str(output_path))
    logger.info("已生成: %s", ruling_name)
    generated_files.append(output_path)


def _insert_extra_paragraphs(ref_para, lines: list[str]):
    """在 ref_para 后面依次插入额外段落（复制 ref_para 格式）。"""
    from docx.oxml.ns import qn
    from copy import deepcopy

    insert_after = ref_para._element
    for line in lines:
        new_para_elem = deepcopy(ref_para._element)
        for r_elem in new_para_elem.findall(qn("w:r")):
            t_elem = r_elem.find(qn("w:t"))
            if t_elem is not None:
                t_elem.text = ""
        first_r = new_para_elem.find(qn("w:r"))
        if first_r is not None:
            t = first_r.find(qn("w:t"))
            if t is not None:
                t.text = line
        insert_after.addnext(new_para_elem)
        insert_after = new_para_elem


# ---------- 保全卷相关 ----------

# 深圳市内鹰眼银行关键词列表
_EAGLE_EYE_BANK_KEYWORDS = [
    "工商银行",
    "农业银行",
    "中国银行",
    "建设银行",
    "交通银行",
    "招商银行",
    "平安银行",
    "农村商业",
    "中信银行",
    "民生银行",
    "华夏银行",
    "兴业银行",
    "上海浦发银行",
    "邮政储蓄",
    "邮储银行",
    "宁波银行",
    "杭州银行",
    "北京银行",
    "江苏银行",
    "包商银行",
    "徽商银行",
    "南洋商业银行",
    "渤海银行",
    "汇丰银行",
    "渣打银行",
]


def _is_eagle_eye_bank(bank_name: str) -> bool:
    """判断开户行是否属于深圳市鹰眼银行。"""
    for keyword in _EAGLE_EYE_BANK_KEYWORDS:
        if keyword in bank_name:
            return True
    return False


def _is_shenzhen_property(clue: dict) -> bool:
    """判断房产是否在深圳市内。"""
    location = clue.get("归属地", "")
    address = clue.get("房产地址", "")
    detail = clue.get("详细内容", "")
    # 优先看归属地，其次看房产地址/详细内容是否包含"深圳"
    if "深圳" in location:
        return True
    if not location and ("深圳" in address or "深圳" in detail):
        return True
    return False


def _is_shenzhen_equity(clue: dict) -> bool:
    """判断股权是否在深圳市内。"""
    location = clue.get("归属地", "")
    detail = clue.get("详细内容", "")
    if "深圳" in location:
        return True
    if not location and "深圳" in detail:
        return True
    return False


def _is_shenzhen_vehicle(clue: dict) -> bool:
    """判断车辆是否在深圳市内。"""
    location = clue.get("归属地", "")
    detail = clue.get("详细内容", "")
    if "深圳" in location:
        return True
    if not location and "深圳" in detail:
        return True
    return False


def _classify_dossier_type(clue: dict) -> str:
    """
    将财产线索分类为 "errand"（外勤）或 "eagle_eye"（鹰眼）。

    分类规则：
    - 支付宝、财付通 → 外勤
    - 车辆：深圳市内 → 鹰眼，市外 → 外勤
    - 股权：深圳市内 → 鹰眼，市外 → 外勤
    - 银行：23家鹰眼银行内 → 鹰眼，其余 → 外勤
    - 房产：深圳市内 → 鹰眼，市外 → 外勤
    - 不支持的类型 → 返回空字符串（跳过）
    """
    clue_type = clue.get("类型", "")
    template_name = _get_template_name(clue_type)
    if template_name is None:
        return ""

    if template_name == "alipay" or template_name == "tenpay":
        return "errand"

    if template_name == "vehicle":
        return "eagle_eye" if _is_shenzhen_vehicle(clue) else "errand"

    if template_name == "equity":
        return "eagle_eye" if _is_shenzhen_equity(clue) else "errand"

    if template_name == "bank":
        bank_name = clue.get("开户银行", "")
        if not bank_name:
            bank_name = _extract_bank_name(clue.get("详细内容", ""))
        return "eagle_eye" if _is_eagle_eye_bank(bank_name) else "errand"

    if template_name == "property":
        return "eagle_eye" if _is_shenzhen_property(clue) else "errand"

    return ""


def _build_dossier_item_text(clue: dict, defendants: list[dict], freeze_amount: str) -> str:
    """根据单条财产线索生成外勤/鹰眼条目文本。"""
    clue_detail = clue.get("详细内容", "")
    clue_type = clue.get("类型", "")
    target = _match_defendant_by_clue(defendants, clue_detail)
    target_name = _get_party_name(target)
    target_id = _get_id_number(target)
    target_id_label = _get_id_label(target)

    template_name = _get_template_name(clue_type)
    if template_name is None:
        return ""

    id_part = f"{target_name}（{target_id_label}{target_id}）" if target_id else target_name

    if template_name == "bank":
        bank_name = clue.get("开户银行", "")
        bank_account = clue.get("银行账号", "")
        if not bank_name:
            bank_name = _extract_bank_name(clue_detail)
        if not bank_account:
            bank_account = _extract_account(clue_detail)
        return f"冻结被告{id_part}在{bank_name}（处）{bank_account}账户的存款人民币{freeze_amount}元，请暂停支付一年。"

    if template_name == "alipay":
        alipay_number = clue.get("支付宝账号", "")
        if not alipay_number:
            alipay_number = _ai_extract_alipay(clue_detail)
        if not re.match(r"^1\d{10}$", alipay_number):
            phone_match = re.search(r"1[3-9]\d{9}", clue_detail)
            if phone_match:
                alipay_number = phone_match.group()
        return f"冻结被告{id_part}名下支付宝账户{alipay_number}内的存款人民币{freeze_amount}元。"

    if template_name == "tenpay":
        tenpay_account = clue.get("微信号", "")
        if not tenpay_account:
            tenpay_account = _ai_extract_tenpay(clue_detail)
        return f"冻结被告{id_part}在财付通账户{tenpay_account}内的存款人民币{freeze_amount}元。"

    if template_name == "property":
        property_address = clue.get("房产地址", "")
        property_cert = clue.get("不动产证号", "")
        if not property_address or not property_cert:
            property_info = _ai_extract_property(clue_detail)
            if not property_address:
                property_address = property_info.get("property_address", "")
            if not property_cert:
                property_cert = property_info.get("property_certificate_number", "")
        return f"查封被告{id_part}名下位于{property_address}的房产100%的份额，不动产权证号：{property_cert}。"

    if template_name == "equity":
        return f"冻结被告{id_part}{clue_detail}"

    if template_name == "vehicle":
        return f"查封被告{id_part}名下{clue_detail}。"

    return ""


def _build_dossier_context(case_data: dict) -> dict:
    """构建保全卷的模板变量（外勤/鹰眼分类）。"""
    case_info = case_data.get("案件基础信息", {})
    full_case_number = case_info.get("案号", "")
    year = _extract_year(full_case_number)
    case_number = _extract_case_number(full_case_number)

    plaintiffs = case_data.get("原告", [])
    defendants = case_data.get("被告", [])
    plaintiff_name = _join_party_names(plaintiffs)
    defendant_name = _join_party_names(defendants)

    preservation_info = case_data.get("保全信息", {})
    freeze_amount = preservation_info.get("申请保全总金额", "")
    clues = preservation_info.get("财产线索", [])

    # 被告段落
    defendant_lines = []
    for d in defendants:
        name = _get_party_name(d)
        if not name:
            continue
        id_number = _get_id_number(d)
        id_label = _get_id_label(d)
        if id_number:
            defendant_lines.append(f"被告：{name}（{id_label}{id_number}）。")
        else:
            defendant_lines.append(f"被告：{name}。")

    # 按外勤/鹰眼分类
    errand_items = []
    eagle_eye_items = []
    for clue in clues:
        dossier_type = _classify_dossier_type(clue)
        if not dossier_type:
            continue
        item_text = _build_dossier_item_text(clue, defendants, freeze_amount)
        if not item_text:
            continue
        if dossier_type == "errand":
            errand_items.append(item_text)
        else:
            eagle_eye_items.append(item_text)

    # 添加中文序号（一、二、三...）
    numbered_errand_items = []
    for idx, item in enumerate(errand_items, 1):
        prefix = f"{_num_to_chinese(idx)}、"
        numbered_errand_items.append(f"{prefix}{item}")

    numbered_eagle_eye_items = []
    for idx, item in enumerate(eagle_eye_items, 1):
        prefix = f"{_num_to_chinese(idx)}、"
        numbered_eagle_eye_items.append(f"{prefix}{item}")

    return {
        "year": year,
        "case_number": case_number,
        "plaintiff_name": plaintiff_name,
        "defendant_name": defendant_name,
        "defendant_lines": defendant_lines,
        "dispute_type": case_info.get("案由", ""),
        "freeze_amount": freeze_amount,
        "errand_items": numbered_errand_items,
        "eagle_eye_items": numbered_eagle_eye_items,
    }


def _remove_paragraph(para):
    """从文档中移除指定段落。"""
    p_elem = para._element
    parent = p_elem.getparent()
    if parent is not None:
        parent.remove(p_elem)


def _generate_dossier(
    case_data: dict, template_path: Path, output_dir: Path, generated_files: list[Path]
):
    """用 python-docx 直接生成保全卷（支持多行被告段落和外勤/鹰眼分类条目）。"""
    from docx import Document

    context = _build_dossier_context(case_data)
    doc = Document(str(template_path))

    # 找到占位段落和标题段落
    defendants_para = None
    errand_title_para = None
    errand_items_para = None
    eagle_eye_title_para = None
    eagle_eye_items_para = None
    # 记录需要处理的空行分隔
    prev_empty_before_errand = None  # 外勤标题前的空行
    prev_empty_before_eagle = None   # 鹰眼标题前的空行

    all_paras = list(doc.paragraphs)
    for idx, p in enumerate(all_paras):
        if "{{ defendants_block }}" in p.text:
            defendants_para = p
        elif p.text == "外勤":
            errand_title_para = p
            # 外勤标题前一个段落是空行
            if idx > 0 and all_paras[idx - 1].text == "":
                prev_empty_before_errand = all_paras[idx - 1]
        elif "{{ errand_items }}" in p.text:
            errand_items_para = p
        elif p.text == "鹰眼":
            eagle_eye_title_para = p
            # 鹰眼标题前一个段落是空行
            if idx > 0 and all_paras[idx - 1].text == "":
                prev_empty_before_eagle = all_paras[idx - 1]
        elif "{{ eagle_eye_items }}" in p.text:
            eagle_eye_items_para = p

    # 插入被告段落
    if defendants_para is not None:
        defendant_lines = context.get("defendant_lines", [])
        if defendant_lines:
            _replace_in_runs(defendants_para, "{{ defendants_block }}", defendant_lines[0])
            _insert_extra_paragraphs(defendants_para, defendant_lines[1:])

    # 处理外勤条目
    errand_items = context.get("errand_items", [])
    if errand_items:
        if errand_items_para is not None:
            _replace_in_runs(errand_items_para, "{{ errand_items }}", errand_items[0])
            _insert_extra_paragraphs(errand_items_para, errand_items[1:])
    else:
        # 无外勤条目，移除标题、占位段落及标题前空行
        if errand_items_para is not None:
            _remove_paragraph(errand_items_para)
        if errand_title_para is not None:
            _remove_paragraph(errand_title_para)
        if prev_empty_before_errand is not None:
            _remove_paragraph(prev_empty_before_errand)

    # 处理鹰眼条目
    eagle_eye_items = context.get("eagle_eye_items", [])
    if eagle_eye_items:
        if eagle_eye_items_para is not None:
            _replace_in_runs(eagle_eye_items_para, "{{ eagle_eye_items }}", eagle_eye_items[0])
            _insert_extra_paragraphs(eagle_eye_items_para, eagle_eye_items[1:])
    else:
        # 无鹰眼条目，移除标题、占位段落及标题前空行
        if eagle_eye_items_para is not None:
            _remove_paragraph(eagle_eye_items_para)
        if eagle_eye_title_para is not None:
            _remove_paragraph(eagle_eye_title_para)
        if prev_empty_before_eagle is not None:
            _remove_paragraph(prev_empty_before_eagle)

    # 替换所有其他 {{ }} 变量
    simple_vars = {
        "{{ year }}": context.get("year", ""),
        "{{ case_number }}": context.get("case_number", ""),
        "{{ plaintiff_name }}": context.get("plaintiff_name", ""),
        "{{ defendant_name }}": context.get("defendant_name", ""),
        "{{ dispute_type }}": context.get("dispute_type", ""),
        "{{ freeze_amount }}": context.get("freeze_amount", ""),
    }
    for p in doc.paragraphs:
        for old, new in simple_vars.items():
            if old in p.text:
                _replace_in_runs(p, old, new)

    case_number = context.get("case_number", "")
    dossier_name = f"{CIVIL_FIRST}{case_number}保全卷.docx"
    output_path = output_dir / dossier_name
    doc.save(str(output_path))
    logger.info("已生成: %s", dossier_name)
    generated_files.append(output_path)


def _get_location_for_naming(clue: dict) -> str:
    """提取地名用于文件命名。'深圳市龙岗区' → '深圳'。"""
    location = clue.get("归属地", "")
    if not location:
        return ""
    idx = location.find("市")
    if idx > 0:
        return location[:idx]
    return location


def _get_base_filename(template_name: str, clue: dict) -> str:
    """根据线索类型和内容生成文件名基础部分（不含扩展名）。"""
    if template_name == "property":
        location = _get_location_for_naming(clue)
        return f"协助执行通知书（查封{location}房产）"
    elif template_name == "bank":
        account = clue.get("银行账号", "") or clue.get("bank账号", "")
        if not account:
            account = _extract_account(clue.get("详细内容", ""))
        last_four = account[-4:] if len(account) >= 4 else account
        return f"协助执行通知书（银行{last_four}）"
    elif template_name == "vehicle":
        location = _get_location_for_naming(clue)
        return f"协助执行通知书（查封{location}车辆）"
    elif template_name == "alipay":
        return "协助执行通知书（支付宝）"
    elif template_name == "tenpay":
        return "协助执行通知书（财付通）"
    elif template_name == "equity":
        return "协助执行通知书（冻结股权）"
    return ""


def generate_docs(case_data: dict, output_dir: Path, civil: bool = False) -> list[Path]:
    """
    根据案件数据生成协助执行通知书和保全裁定书。

    每条财产线索生成一份通知书，每案生成一份裁定书。
    文件命名规则：
    - 房产：协助执行通知书（查封xx房产）.docx，同地多套加中文序号
    - 银行：协助执行通知书（银行xxxx）.docx，尾号相同加 -1, -2...
    - 车辆：协助执行通知书（查封xx车辆）.docx，同地多辆加中文序号
    - 支付宝/财付通/股权：协助执行通知书（xxx）.docx，多个加中文序号

    Args:
        case_data: AI提取的结构化案件数据
        output_dir: 输出目录
        civil: 民事模式，只生成外勤类通知书（跳过鹰眼和支付宝）

    Returns:
        生成的文件路径列表
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    generated_files = []

    # 生成保全裁定书（每案一份）
    ruling_template = TEMPLATES_DIR / "ruling.docx"
    if ruling_template.exists():
        _generate_ruling(case_data, ruling_template, output_dir, generated_files)

    # 生成保全卷（每案一份）
    dossier_template = TEMPLATES_DIR / "dossier.docx"
    if dossier_template.exists():
        _generate_dossier(case_data, dossier_template, output_dir, generated_files)

    preservation_info = case_data.get("保全信息", {})
    clues = preservation_info.get("财产线索", [])

    if not clues:
        logger.warning("没有找到财产线索，跳过文档生成")
        return []

    # 第一轮：计算每条线索的模板、基础文件名和分组数量
    clue_info = []
    group_counts = {}  # group_key → count

    for clue in clues:
        clue_type = clue.get("类型", "")
        template_name = _get_template_name(clue_type)
        if template_name is None:
            clue_info.append((clue, None, None, None))
            continue

        # 民事模式：跳过鹰眼类和支付宝
        if civil:
            dossier_type = _classify_dossier_type(clue)
            if template_name == "alipay" or dossier_type == "eagle_eye":
                logger.info("民事模式跳过: %s（%s）", clue_type, "支付宝" if template_name == "alipay" else "鹰眼")
                continue

        base = _get_base_filename(template_name, clue)

        # 银行按基础名分组（尾号相同时重名），房产/车辆按基础名分组，其余按类型分组
        if template_name in ("bank", "property", "vehicle"):
            group_key = base
        else:
            group_key = template_name

        group_counts[group_key] = group_counts.get(group_key, 0) + 1
        clue_info.append((clue, template_name, base, group_key))

    # 第二轮：生成文件，处理重名
    group_seq = {}

    for clue, template_name, base, group_key in clue_info:
        if template_name is None:
            logger.info("跳过不支持的财产线索类型: %s", clue.get("类型", ""))
            continue

        template_path = TEMPLATES_DIR / f"{template_name}.docx"
        if not template_path.exists():
            logger.warning("模板不存在: %s，跳过", template_path)
            continue

        group_seq[group_key] = group_seq.get(group_key, 0) + 1
        seq = group_seq[group_key]
        total = group_counts[group_key]

        if template_name == "bank":
            # 银行：第一个无后缀，重复的加 -1, -2...
            if seq == 1:
                output_name = f"{base}.docx"
            else:
                output_name = f"{base}-{seq - 1}.docx"
        elif total > 1:
            # 同组多个：全部加中文序号（插入在末尾 ）之前）
            suffix = _num_to_chinese(seq)
            output_name = f"{base[:-1]}{suffix}）.docx"
        else:
            output_name = f"{base}.docx"

        context = _build_context(case_data, clue)
        output_path = output_dir / output_name

        tpl = DocxTemplate(str(template_path))
        tpl.render(context)
        tpl.save(str(output_path))

        logger.info("已生成: %s", output_name)
        generated_files.append(output_path)

    return generated_files
